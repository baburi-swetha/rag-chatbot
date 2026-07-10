import os
import pandas as pd

from docx import Document as DocxDocument

from langchain_core.documents import Document

from docling.document_converter import DocumentConverter

converter = DocumentConverter()


def load_document(file_path):

    extension = os.path.splitext(file_path)[1].lower()


    if extension == ".pdf":

        result = converter.convert(file_path)

    
        text = result.document.export_to_markdown()

        return [
            Document(
                page_content=text,
                metadata={}
            )
        ]

   
    elif extension == ".xlsx":

        df = pd.read_excel(file_path)

        text = df.to_string(index=False)

        return [
            Document(
                page_content=text,
                metadata={}
            )
        ]

  
    elif extension == ".docx":

        doc = DocxDocument(file_path)

        text = "\n".join(
            para.text
            for para in doc.paragraphs
            if para.text.strip()
        )

        return [
            Document(
                page_content=text,
                metadata={}
            )
        ]

    elif extension == ".txt":

        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()

        return [
            Document(
                page_content=text,
                metadata={}
            )
        ]

    else:

        raise ValueError(
            f"Unsupported file format: {extension}"
        )
